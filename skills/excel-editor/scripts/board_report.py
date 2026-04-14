"""
Board Report Generator — DataBaseSearch.xlsm
Usage:
  python board_report.py --sheet "Summary" --output "excel,html" --extra "top 5 customers"
  python board_report.py --sheet "Unit 5" --output "pdf,word"
  python board_report.py --folder "C:/path" --workbook "file.xlsm" --sheet "Summary" --output "all"
"""

import sys
import json
import argparse
import openpyxl
from openpyxl.styles import PatternFill, Font, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from datetime import date
from collections import Counter

# ── Defaults ─────────────────────────────────────────────
DEFAULT_FOLDER   = "C:/Users/Darren/OneDrive - SGM Windows/Desktop/Folders/Claude/Data sets"
DEFAULT_WORKBOOK = "DataBaseSearch.xlsm"

# ── Argument parsing ──────────────────────────────────────
parser = argparse.ArgumentParser()
parser.add_argument("--folder",   default=None)
parser.add_argument("--workbook", default=None)
parser.add_argument("--sheet",    default=None)
parser.add_argument("--output",   default=None, help="Comma-separated: excel,html,word,pdf,all")
parser.add_argument("--extra",    default=None, help="Specific info to highlight in the report")
parser.add_argument("--dates",    default=None, help="Search for dates: yes/no")
args = parser.parse_args()

folder   = args.folder   or DEFAULT_FOLDER
workbook = args.workbook or DEFAULT_WORKBOOK
wb_path  = folder.rstrip("/") + "/" + workbook

# ── Load workbook ─────────────────────────────────────────
try:
    wb = openpyxl.load_workbook(wb_path, keep_vba=True, data_only=True)
except FileNotFoundError:
    print(f"ERROR: File not found: {wb_path}")
    sys.exit(1)

available = wb.sheetnames
skip = {"Board Report", "Charts", "Update", "Data"}
data_sheets = [s for s in available if s not in skip]

# ── Sheet selection ───────────────────────────────────────
if args.sheet:
    sheet_name = args.sheet
    if sheet_name not in available:
        print(f"ERROR: Sheet '{sheet_name}' not found.")
        print("Available sheets:", ", ".join(available))
        sys.exit(1)
else:
    print("\nAvailable sheets:")
    for i, s in enumerate(data_sheets, 1):
        print(f"  {i}. {s}")
    print("  A. All data sheets")
    choice = input("\nEnter sheet name or number (default: Summary): ").strip()
    if not choice or choice.upper() == "SUMMARY":
        sheet_name = "Summary"
    elif choice.upper() == "A":
        sheet_name = "ALL"
    elif choice.isdigit() and 1 <= int(choice) <= len(data_sheets):
        sheet_name = data_sheets[int(choice) - 1]
    elif choice in available:
        sheet_name = choice
    else:
        print(f"Invalid choice '{choice}', defaulting to Summary.")
        sheet_name = "Summary"

sheets_to_process = data_sheets if sheet_name == "ALL" else [sheet_name]

# ── Output format selection ───────────────────────────────
if args.output:
    raw = args.output.lower()
    if raw == "all":
        outputs = {"excel", "html", "word", "pdf"}
    else:
        outputs = {o.strip() for o in raw.split(",")}
else:
    print("\nOutput format(s):")
    print("  1. Excel")
    print("  2. HTML")
    print("  3. Word")
    print("  4. PDF")
    print("  A. All")
    fmt_choice = input("\nEnter number(s) separated by commas (default: 1,2): ").strip()
    fmt_map = {"1":"excel","2":"html","3":"word","4":"pdf"}
    if not fmt_choice or fmt_choice.upper() == "A":
        outputs = {"excel","html","word","pdf"} if fmt_choice.upper() == "A" else {"excel","html"}
    else:
        outputs = {fmt_map[c.strip()] for c in fmt_choice.split(",") if c.strip() in fmt_map}
    if not outputs:
        outputs = {"excel","html"}

# ── Extra info request ────────────────────────────────────
extra_info = args.extra or ""
if not args.extra:
    extra_input = input("\nAny specific information to highlight? (press Enter to skip): ").strip()
    extra_info = extra_input

# ── Date search option ────────────────────────────────────
if args.dates:
    search_dates = args.dates.lower() in ("yes", "y", "true", "1")
else:
    date_ans = input("\nSearch for dates in the data? (Y/N, default N): ").strip().lower()
    search_dates = date_ans in ("y", "yes")

print(f"\nProcessing: {', '.join(sheets_to_process)} | Outputs: {', '.join(sorted(outputs))}")

# ── Read & analyse sheet(s) ───────────────────────────────
def read_sheet(wb, name):
    ws = wb[name]
    rows = []
    for row in ws.iter_rows(values_only=True):
        rows.append(list(row))
    return rows

def analyse(rows, sheet_name):
    if not rows or not rows[0]:
        return None
    header = rows[0]

    # Map column names to indices (case-insensitive strip)
    col = {}
    for i, h in enumerate(header):
        if h:
            col[str(h).strip().lower()] = i

    def idx(candidates):
        for c in candidates:
            if c in col:
                return col[c]
        return None

    i_partcode  = idx(["part code"])
    i_desc      = idx(["part description"])
    i_qty       = idx(["quantity"])
    i_customer  = idx(["customer"])
    i_job       = idx(["job number"])
    i_delivery  = idx(["delivery method"])
    i_received  = idx(["parts received?"])
    i_dpd       = idx(["dpd code"])
    i_paperwork = idx(["paperwork?"])
    i_cost_item = idx(["cost per item ","cost per item"])
    i_total     = idx(["total cost ","total cost"])
    i_source    = idx(["source tab"])

    data_rows = [r for r in rows[1:] if i_partcode is not None and len(r) > i_partcode and r[i_partcode]]

    total_cost  = sum(r[i_total] for r in data_rows if i_total is not None and r[i_total] and isinstance(r[i_total], (int,float)))
    total_qty   = sum(r[i_qty]   for r in data_rows if i_qty   is not None and r[i_qty]   and isinstance(r[i_qty],   (int,float)))

    # Group by source tab if present, else use sheet name
    by_unit = {}
    by_customer = {}
    for r in data_rows:
        unit = r[i_source] if i_source is not None and r[i_source] else sheet_name
        cost = r[i_total]  if i_total  is not None and isinstance(r[i_total], (int,float)) else 0
        by_unit[unit] = by_unit.get(unit, {"items": 0, "cost": 0.0})
        by_unit[unit]["items"] += 1
        by_unit[unit]["cost"]  += cost

        if i_customer is not None:
            cust = r[i_customer] if r[i_customer] else "Unknown"
            by_customer[cust] = by_customer.get(cust, 0) + cost

    missing_job      = [r for r in data_rows if i_job       is not None and not r[i_job]]
    missing_delivery = [r for r in data_rows if i_delivery  is not None and not r[i_delivery]]
    missing_received = [r for r in data_rows if i_received  is not None and not r[i_received]]
    missing_paper    = [r for r in data_rows if i_paperwork is not None and (not r[i_paperwork] or str(r[i_paperwork]).strip() != "YES")]
    dpd_unconfirmed  = [r for r in data_rows if i_dpd       is not None and r[i_dpd] and i_received is not None and not r[i_received]]

    part_codes = [str(r[i_partcode]).strip().upper() for r in data_rows]
    dupes = {k: v for k, v in Counter(part_codes).items() if v > 1}

    top_customers = sorted(by_customer.items(), key=lambda x: -x[1])[:10]

    return {
        "sheet": sheet_name,
        "line_items": len(data_rows),
        "total_cost": total_cost,
        "total_qty": total_qty,
        "by_unit": by_unit,
        "top_customers": top_customers,
        "missing_job": len(missing_job),
        "missing_delivery": len(missing_delivery),
        "missing_received": len(missing_received),
        "missing_paper": len(missing_paper),
        "dpd_unconfirmed": len(dpd_unconfirmed),
        "duplicate_part_codes": len(dupes),
    }

# ── Date scanning ────────────────────────────────────────
import re
from datetime import datetime as dt

def scan_dates(wb, sheet_names):
    """Scan sheets for date values (cell dates + DD.MM.YY patterns in text)."""
    found = []
    date_pattern = re.compile(r'\b(\d{1,2})[.\-/](\d{1,2})[.\-/](\d{2,4})\b')
    for sname in sheet_names:
        ws_scan = wb[sname]
        for row in ws_scan.iter_rows(values_only=True):
            for cell in row:
                if cell is None:
                    continue
                # Native datetime cells
                if hasattr(cell, 'year'):
                    try:
                        found.append(dt(cell.year, cell.month, cell.day))
                    except Exception:
                        pass
                # Date patterns inside strings (e.g. "Parts List - 02.03.26.xlsx")
                if isinstance(cell, str):
                    for m in date_pattern.finditer(cell):
                        d_val, mo_val, y_val = int(m.group(1)), int(m.group(2)), int(m.group(3))
                        if y_val < 100:
                            y_val += 2000
                        try:
                            if 1 <= d_val <= 31 and 1 <= mo_val <= 12:
                                found.append(dt(y_val, mo_val, d_val))
                        except Exception:
                            pass
    return found

date_range_str = ""
if search_dates:
    all_dates = scan_dates(wb, sheets_to_process)
    if all_dates:
        min_d = min(all_dates).strftime("%d/%m/%Y")
        max_d = max(all_dates).strftime("%d/%m/%Y")
        date_range_str = f"{min_d} — {max_d}" if min_d != max_d else min_d
        print(f"Dates found: {date_range_str}  ({len(all_dates)} date values)")
    else:
        print("No dates found in selected sheet(s).")

# Collect analysis across selected sheets
all_data = []
for sname in sheets_to_process:
    rows = read_sheet(wb, sname)
    result = analyse(rows, sname)
    if result:
        all_data.append(result)

if not all_data:
    print("No data found in selected sheet(s).")
    sys.exit(1)

# Aggregate if multiple sheets
if len(all_data) > 1:
    agg = {
        "sheet": "All Sheets",
        "line_items": sum(d["line_items"] for d in all_data),
        "total_cost": sum(d["total_cost"] for d in all_data),
        "total_qty": sum(d["total_qty"] for d in all_data),
        "by_unit": {},
        "top_customers": {},
        "missing_job": sum(d["missing_job"] for d in all_data),
        "missing_delivery": sum(d["missing_delivery"] for d in all_data),
        "missing_received": sum(d["missing_received"] for d in all_data),
        "missing_paper": sum(d["missing_paper"] for d in all_data),
        "dpd_unconfirmed": sum(d["dpd_unconfirmed"] for d in all_data),
        "duplicate_part_codes": sum(d["duplicate_part_codes"] for d in all_data),
    }
    for d in all_data:
        for unit, v in d["by_unit"].items():
            agg["by_unit"][unit] = agg["by_unit"].get(unit, {"items":0,"cost":0.0})
            agg["by_unit"][unit]["items"] += v["items"]
            agg["by_unit"][unit]["cost"]  += v["cost"]
        for cust, cost in d["top_customers"]:
            agg["top_customers"][cust] = agg["top_customers"].get(cust, 0) + cost
    agg["top_customers"] = sorted(agg["top_customers"].items(), key=lambda x: -x[1])[:10]
    stats = agg
else:
    stats = all_data[0]

# ── Style helpers ─────────────────────────────────────────
def fill(hex): return PatternFill("solid", fgColor=hex)
def fnt(bold=False, size=11, color="000000"): return Font(bold=bold, size=size, color=color, name="Calibri")
def aln(h="left", v="center", wrap=False): return Alignment(horizontal=h, vertical=v, wrap_text=wrap)
def bdr():
    s = Side(style="thin")
    return Border(left=s, right=s, top=s, bottom=s)

DARK_BLUE   = fill("1F3864")
MID_BLUE    = fill("2E5FA3")
LIGHT_BLUE  = fill("D6E4F7")
RED_FILL    = fill("C00000")
ORANGE_FILL = fill("FF8C00")
YELLOW_FILL = fill("FFD700")
GREEN_FILL  = fill("375623")
GREY_FILL   = fill("F2F2F2")
WHITE_FILL  = fill("FFFFFF")
THIN        = bdr()

def sc(ws, row, col, value, f=None, fi=None, a=None):
    c = ws.cell(row=row, column=col, value=value)
    if f:  c.font = f
    if fi: c.fill = fi
    if a:  c.alignment = a
    c.border = THIN
    return c

def ms(ws, row, c1, c2, value, f=None, fi=None, a=None):
    ws.merge_cells(start_row=row, start_column=c1, end_row=row, end_column=c2)
    c = ws.cell(row=row, column=c1, value=value)
    if f:  c.font = f
    if fi: c.fill = fi
    if a:  c.alignment = a
    return c

# ── Build Board Report sheet ──────────────────────────────
if "Board Report" in wb.sheetnames:
    del wb["Board Report"]
ws = wb.create_sheet("Board Report", 0)

for col, w in zip(range(1, 9), [2, 22, 14, 14, 12, 14, 28, 2]):
    ws.column_dimensions[get_column_letter(col)].width = w

today  = date.today().strftime("%d/%m/%Y")
n      = stats["line_items"]
tc     = stats["total_cost"]
qty    = stats["total_qty"]
src_label = stats["sheet"]

# ROW 1 — Title
ws.row_dimensions[1].height = 40
title_text = "PARTS DATABASE  -  SUMMARY" + (f"  |  {date_range_str}" if date_range_str else "")
ms(ws, 1, 2, 6, title_text, fnt(True,14,"FFFFFF"), DARK_BLUE, aln("left","center"))
sc(ws, 1, 7, "Date: " + today, fnt(False,9,"FFFFFF"), DARK_BLUE, aln("right","center"))
sc(ws, 1, 1, None, fi=DARK_BLUE); sc(ws, 1, 8, None, fi=DARK_BLUE)

# ROW 2 — Subtitle
ws.row_dimensions[2].height = 18
ms(ws, 2, 1, 8, f"Source: {workbook}  |  Sheet: {src_label}  |  {n} Line Items",
   fnt(False,9,"FFFFFF"), MID_BLUE, aln("left","center"))

ws.row_dimensions[3].height = 8

# ROW 4 — KPI header
ws.row_dimensions[4].height = 22
ms(ws, 4, 1, 8, "  KPI OVERVIEW", fnt(True,10,"FFFFFF"), DARK_BLUE, aln("left","center"))

# ROW 5 — KPI boxes
ws.row_dimensions[5].height = 45
active_units = len(stats["by_unit"])
kpis = [
    ("TOTAL COST\n" + ("£{:,.2f}".format(tc) if tc else "N/A"), 2),
    ("TOTAL QTY\n" + str(int(qty) if qty else "N/A"),             3),
    ("LINE ITEMS\n" + str(n),                                      4),
    ("UNITS ACTIVE\n" + str(active_units),                        5),
]
for label, col in kpis:
    sc(ws, 5, col, label, fnt(True,10,"1F3864"), LIGHT_BLUE, aln("center","center",True))

ws.row_dimensions[6].height = 8

# ROW 7 — Cost by Unit header
ws.row_dimensions[7].height = 22
ms(ws, 7, 1, 8, "  COST & VOLUME BY UNIT", fnt(True,10,"FFFFFF"), DARK_BLUE, aln("left","center"))

# ROW 8 — Table headers
ws.row_dimensions[8].height = 22
for col, h in zip([2,3,4,5,6], ["Unit","Line Items","Total Cost","% of Total","Avg Cost/Item"]):
    sc(ws, 8, col, h, fnt(True,10,"FFFFFF"), MID_BLUE, aln("center","center"))

# ROW 9+ — Unit rows (sorted by cost desc)
unit_rows = sorted(stats["by_unit"].items(), key=lambda x: -x[1]["cost"])
for r, (unit, v) in enumerate(unit_rows, 9):
    ws.row_dimensions[r].height = 18
    rf   = GREY_FILL if r % 2 == 0 else WHITE_FILL
    cost = v["cost"]; items = v["items"]
    sc(ws, r, 2, unit,                              fnt(False,10), rf, aln("left","center"))
    sc(ws, r, 3, items,                             fnt(False,10), rf, aln("center","center"))
    sc(ws, r, 4, "£{:,.2f}".format(cost),          fnt(False,10), rf, aln("center","center"))
    sc(ws, r, 5, "{:.1f}%".format(cost/tc*100 if tc else 0), fnt(False,10), rf, aln("center","center"))
    sc(ws, r, 6, "£{:.2f}".format(cost/items if items else 0), fnt(False,10), rf, aln("center","center"))

# Total row
tr = 9 + len(unit_rows)
ws.row_dimensions[tr].height = 20
for col, val in zip([2,3,4,5,6], ["TOTAL", str(n), "£{:,.2f}".format(tc), "100%", "£{:.2f}".format(tc/n if n else 0)]):
    sc(ws, tr, col, val, fnt(True,10,"FFFFFF"), DARK_BLUE, aln("center","center"))
ws.cell(row=tr, column=2).alignment = aln("left","center")

spacer1 = tr + 1
ws.row_dimensions[spacer1].height = 8

# Top Customers
cust_hdr = spacer1 + 1
ws.row_dimensions[cust_hdr].height = 22
ms(ws, cust_hdr, 1, 8, "  TOP CUSTOMERS BY COST", fnt(True,10,"FFFFFF"), DARK_BLUE, aln("left","center"))

cust_col_hdr = cust_hdr + 1
ws.row_dimensions[cust_col_hdr].height = 22
for col, h in zip([2,3,4], ["Customer","Total Cost","% of Total"]):
    sc(ws, cust_col_hdr, col, h, fnt(True,10,"FFFFFF"), MID_BLUE, aln("center","center"))

for r, (cust, cost) in enumerate(stats["top_customers"], cust_col_hdr + 1):
    ws.row_dimensions[r].height = 18
    rf = GREY_FILL if r % 2 == 0 else WHITE_FILL
    sc(ws, r, 2, cust,                                     fnt(False,10), rf, aln("left","center"))
    sc(ws, r, 3, "£{:,.2f}".format(cost),                 fnt(False,10), rf, aln("center","center"))
    sc(ws, r, 4, "{:.1f}%".format(cost/tc*100 if tc else 0), fnt(False,10), rf, aln("center","center"))

spacer2 = cust_col_hdr + len(stats["top_customers"]) + 1
ws.row_dimensions[spacer2].height = 8

# Anomalies
anom_hdr = spacer2 + 1
ws.row_dimensions[anom_hdr].height = 22
ms(ws, anom_hdr, 1, 8, "  ANOMALIES & DATA QUALITY FLAGS", fnt(True,10,"FFFFFF"), RED_FILL, aln("left","center"))

anom_col_hdr = anom_hdr + 1
ws.row_dimensions[anom_col_hdr].height = 22
for col, h in zip([2,3,4,5,6], ["Flag","Count","% of Total","Severity","Action Required"]):
    sc(ws, anom_col_hdr, col, h, fnt(True,10,"FFFFFF"), MID_BLUE, aln("center","center"))

def sev(label):
    if label == "CRITICAL": return fill("C00000"), fnt(True,9,"FFFFFF")
    if label == "HIGH":     return ORANGE_FILL,    fnt(True,9,"FFFFFF")
    if label == "MEDIUM":   return YELLOW_FILL,    fnt(True,9,"000000")
    if label == "RESOLVED": return GREEN_FILL,     fnt(True,9,"FFFFFF")
    return WHITE_FILL, fnt(False,9)

anomalies = []
if stats["missing_received"] > 0:
    anomalies.append(("Parts received — not confirmed", f"{stats['missing_received']}/{n}", f"{stats['missing_received']/n*100:.0f}%", "CRITICAL", "Verify receipt for all items"))
if stats["missing_paper"] > 0:
    anomalies.append(("Missing paperwork",              f"{stats['missing_paper']}/{n}",    f"{stats['missing_paper']/n*100:.0f}%",    "HIGH",     "Obtain / log paperwork"))
if stats["missing_job"] > 0:
    anomalies.append(("Missing job number",             f"{stats['missing_job']}/{n}",      f"{stats['missing_job']/n*100:.0f}%",      "HIGH",     "Assign job numbers"))
if stats["missing_delivery"] > 0:
    anomalies.append(("Missing delivery method",        f"{stats['missing_delivery']}/{n}", f"{stats['missing_delivery']/n*100:.0f}%", "MEDIUM",   "Assign delivery route"))
if stats["dpd_unconfirmed"] > 0:
    anomalies.append(("DPD coded — receipt unconfirmed",f"{stats['dpd_unconfirmed']}/{n}",  f"{stats['dpd_unconfirmed']/n*100:.0f}%",  "HIGH",     "Chase DPD confirmations"))
if stats["duplicate_part_codes"] > 0:
    anomalies.append(("Duplicate part codes",           f"{stats['duplicate_part_codes']} codes", "—", "MEDIUM", "Audit for double-ordering"))
if not anomalies:
    anomalies.append(("No anomalies detected", "—", "—", "RESOLVED", "No action required"))

for r, (flag, count, pct, severity, action) in enumerate(anomalies, anom_col_hdr + 1):
    ws.row_dimensions[r].height = 18
    rf = GREY_FILL if r % 2 == 0 else WHITE_FILL
    sev_fill, sev_fnt = sev(severity)
    sc(ws, r, 2, flag,     fnt(False,10), rf,       aln("left","center"))
    sc(ws, r, 3, count,    fnt(False,10), rf,       aln("center","center"))
    sc(ws, r, 4, pct,      fnt(False,10), rf,       aln("center","center"))
    sc(ws, r, 5, severity, sev_fnt,       sev_fill, aln("center","center"))
    sc(ws, r, 6, action,   fnt(False,10), rf,       aln("left","center",True))

spacer3 = anom_col_hdr + len(anomalies) + 1
ws.row_dimensions[spacer3].height = 8

# Legend
leg_hdr = spacer3 + 1
ws.row_dimensions[leg_hdr].height = 22
ms(ws, leg_hdr, 1, 8, "  SUMMARY TAB COLOUR KEY", fnt(True,10,"FFFFFF"), DARK_BLUE, aln("left","center"))

legend = [
    (fill("FFCCCC"), "Light Red  -  Parts received: unconfirmed  (Col I)"),
    (ORANGE_FILL,    "Orange  -  Missing job number / paperwork / DPD unconfirmed  (Cols G, J, K)"),
    (YELLOW_FILL,    "Yellow  -  Missing delivery method  (Col H)"),
]
for r, (leg_fill, label) in enumerate(legend, leg_hdr + 1):
    ws.row_dimensions[r].height = 18
    ws.cell(row=r, column=2).fill = leg_fill
    ws.cell(row=r, column=2).border = THIN
    ms(ws, r, 3, 7, label, fnt(False,10), WHITE_FILL, aln("left","center"))
    for c in range(3, 8):
        ws.cell(row=r, column=c).border = THIN

ws.sheet_properties.tabColor = "1F3864"

# Only save Excel if requested
if "excel" in outputs:
    wb.save(wb_path)
    print(f"Excel: {wb_path}")
else:
    wb.save(wb_path)  # always save to preserve VBA/structure

# ── HTML Output ───────────────────────────────────────────
def sev_html_color(label):
    return {"CRITICAL":"#C00000","HIGH":"#FF8C00","MEDIUM":"#FFD700","RESOLVED":"#375623"}.get(label,"#999")

def sev_html_text(label):
    return "#ffffff" if label in ("CRITICAL","HIGH","RESOLVED") else "#000000"

unit_rows_sorted = sorted(stats["by_unit"].items(), key=lambda x: -x[1]["cost"])

unit_rows_html = ""
for unit, v in unit_rows_sorted:
    unit_rows_html += f"""
        <tr>
          <td>{unit}</td>
          <td style="text-align:center">{v['items']}</td>
          <td style="text-align:center">&pound;{v['cost']:,.2f}</td>
          <td style="text-align:center">{v['cost']/tc*100:.1f}%</td>
          <td style="text-align:center">&pound;{v['cost']/v['items']:.2f}</td>
        </tr>"""

cust_rows_html = ""
for cust, cost in stats["top_customers"]:
    cust_rows_html += f"""
        <tr>
          <td>{cust}</td>
          <td style="text-align:center">&pound;{cost:,.2f}</td>
          <td style="text-align:center">{cost/tc*100:.1f}%</td>
        </tr>"""

anom_rows_html = ""
for flag, count, pct, severity, action in anomalies:
    bg = sev_html_color(severity)
    fg = sev_html_text(severity)
    anom_rows_html += f"""
        <tr>
          <td>{flag}</td>
          <td style="text-align:center">{count}</td>
          <td style="text-align:center">{pct}</td>
          <td style="text-align:center;background:{bg};color:{fg};font-weight:bold;border-radius:4px">{severity}</td>
          <td>{action}</td>
        </tr>"""

html = f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>Board Report &mdash; {src_label}</title>
<style>
  * {{ box-sizing: border-box; margin: 0; padding: 0; }}
  body {{ font-family: Calibri, Arial, sans-serif; background: #f0f4f8; color: #222; }}
  .wrapper {{ max-width: 960px; margin: 32px auto; padding: 0 16px; }}

  /* Header */
  .header {{ background: #1F3864; color: #fff; padding: 24px 28px 16px; border-radius: 8px 8px 0 0; }}
  .header h1 {{ font-size: 26px; letter-spacing: 1px; }}
  .header .meta {{ font-size: 12px; color: #aec6e8; margin-top: 6px; }}
  .subbar {{ background: #2E5FA3; color: #cde; font-size: 12px; padding: 8px 28px; margin-bottom: 24px; border-radius: 0 0 8px 8px; }}

  /* KPI boxes */
  .kpi-grid {{ display: grid; grid-template-columns: repeat(4, 1fr); gap: 12px; margin-bottom: 28px; }}
  .kpi {{ background: #D6E4F7; border-radius: 8px; padding: 16px 12px; text-align: center; border: 1px solid #b3cee8; }}
  .kpi .label {{ font-size: 11px; font-weight: bold; color: #1F3864; text-transform: uppercase; letter-spacing: .5px; }}
  .kpi .value {{ font-size: 22px; font-weight: bold; color: #1F3864; margin-top: 6px; }}

  /* Section headers */
  .section-hdr {{ background: #1F3864; color: #fff; font-size: 13px; font-weight: bold;
                  padding: 10px 16px; border-radius: 6px 6px 0 0; letter-spacing: .5px; margin-top: 24px; }}
  .section-hdr.red {{ background: #C00000; }}

  /* Tables */
  table {{ width: 100%; border-collapse: collapse; font-size: 13px; background: #fff;
           border-radius: 0 0 6px 6px; overflow: hidden; box-shadow: 0 2px 6px rgba(0,0,0,.08); }}
  thead th {{ background: #2E5FA3; color: #fff; padding: 9px 12px; text-align: center; font-size: 12px; }}
  thead th:first-child {{ text-align: left; }}
  tbody tr:nth-child(even) {{ background: #f2f2f2; }}
  tbody tr:hover {{ background: #e8f0fb; }}
  tbody td {{ padding: 8px 12px; border-bottom: 1px solid #e0e0e0; }}
  tfoot td {{ background: #1F3864; color: #fff; font-weight: bold; padding: 9px 12px; text-align: center; }}
  tfoot td:first-child {{ text-align: left; }}

  /* Legend */
  .legend {{ display: flex; gap: 16px; flex-wrap: wrap; margin-top: 24px; font-size: 12px; }}
  .legend-item {{ display: flex; align-items: center; gap: 8px; }}
  .legend-swatch {{ width: 20px; height: 20px; border-radius: 4px; border: 1px solid #ccc; flex-shrink:0; }}

  .footer {{ text-align: center; font-size: 11px; color: #999; margin: 32px 0 16px; }}
</style>
</head>
<body>
<div class="wrapper">

  <div class="header">
    <h1>PARTS DATABASE &mdash; SUMMARY</h1>
    <div class="meta">Date: {today} &nbsp;|&nbsp; Source: {workbook} &nbsp;|&nbsp; Sheet: {src_label}{(" &nbsp;|&nbsp; Period: " + date_range_str) if date_range_str else ""}</div>
  </div>
  <div class="subbar">{n} line items analysed</div>

  <!-- KPIs -->
  <div class="kpi-grid">
    <div class="kpi"><div class="label">Total Cost</div><div class="value">&pound;{tc:,.2f}</div></div>
    <div class="kpi"><div class="label">Total Qty</div><div class="value">{int(qty)}</div></div>
    <div class="kpi"><div class="label">Line Items</div><div class="value">{n}</div></div>
    <div class="kpi"><div class="label">Units Active</div><div class="value">{len(stats['by_unit'])}</div></div>
  </div>

  <!-- Cost by Unit -->
  <div class="section-hdr">COST &amp; VOLUME BY UNIT</div>
  <table>
    <thead><tr><th>Unit</th><th>Line Items</th><th>Total Cost</th><th>% of Total</th><th>Avg Cost/Item</th></tr></thead>
    <tbody>{unit_rows_html}</tbody>
    <tfoot><tr><td>TOTAL</td><td style="text-align:center">{n}</td><td style="text-align:center">&pound;{tc:,.2f}</td><td style="text-align:center">100%</td><td style="text-align:center">&pound;{tc/n:.2f}</td></tr></tfoot>
  </table>

  <!-- Top Customers -->
  <div class="section-hdr">TOP CUSTOMERS BY COST</div>
  <table>
    <thead><tr><th>Customer</th><th>Total Cost</th><th>% of Total</th></tr></thead>
    <tbody>{cust_rows_html}</tbody>
  </table>

  <!-- Anomalies -->
  <div class="section-hdr red">ANOMALIES &amp; DATA QUALITY FLAGS</div>
  <table>
    <thead><tr><th>Flag</th><th>Count</th><th>% of Total</th><th>Severity</th><th>Action Required</th></tr></thead>
    <tbody>{anom_rows_html}</tbody>
  </table>

  <!-- Legend -->
  <div class="legend">
    <div class="legend-item"><div class="legend-swatch" style="background:#FFCCCC"></div> Parts received unconfirmed (Col I)</div>
    <div class="legend-item"><div class="legend-swatch" style="background:#FF8C00"></div> Missing job / paperwork / DPD (Cols G, J, K)</div>
    <div class="legend-item"><div class="legend-swatch" style="background:#FFD700"></div> Missing delivery method (Col H)</div>
  </div>

  <div class="footer">Generated by Claude Code &mdash; {today}</div>
</div>
</body>
</html>"""

import subprocess
import os

base_name = folder.rstrip("/") + "/Board_Report_" + src_label.replace(" ", "_")
html_path = base_name + ".html"

# Always build HTML (needed for PDF conversion too)
with open(html_path, "w", encoding="utf-8") as f:
    f.write(html)

if "html" in outputs:
    print(f"HTML:  {html_path}")
    subprocess.Popen(["start", html_path], shell=True)

# ── Word Output ───────────────────────────────────────────
if "word" in outputs:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def set_cell_bg(cell, hex_color):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tcPr.append(shd)

    def add_table_row(table, cells, bold=False, bg=None, center=False):
        row = table.add_row()
        for i, text in enumerate(cells):
            cell = row.cells[i]
            cell.text = str(text) if text is not None else ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center or i > 0 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.runs[0] if p.runs else p.add_run(str(text) if text is not None else "")
            run.font.bold = bold
            run.font.size = Pt(10)
            if bg:
                set_cell_bg(cell, bg)
                if bg in ("1F3864","2E5FA3","C00000","375623","FF8C00"):
                    run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        return row

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # Margins
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # Title
    title = doc.add_heading("PARTS DATABASE  \u2014  SUMMARY", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.runs[0].font.color.rgb = RGBColor(0x1F,0x38,0x64)

    meta_line = f"Date: {today}  |  Source: {workbook}  |  Sheet: {src_label}  |  {n} line items"
    if date_range_str:
        meta_line += f"  |  Period: {date_range_str}"
    doc.add_paragraph(meta_line)

    if extra_info:
        p = doc.add_paragraph()
        r = p.add_run(f"Focus: {extra_info}")
        r.font.bold = True
        r.font.color.rgb = RGBColor(0xC0,0x00,0x00)

    # KPI section
    doc.add_heading("KPI OVERVIEW", 2)
    kpi_table = doc.add_table(rows=1, cols=4)
    kpi_table.style = "Table Grid"
    hdr = kpi_table.rows[0].cells
    for i, (label, val) in enumerate([
        ("TOTAL COST", f"\xa3{tc:,.2f}"),
        ("TOTAL QTY",  str(int(qty))),
        ("LINE ITEMS", str(n)),
        ("UNITS ACTIVE", str(len(stats["by_unit"]))),
    ]):
        hdr[i].text = f"{label}\n{val}"
        set_cell_bg(hdr[i], "D6E4F7")
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x1F,0x38,0x64)

    # Cost by Unit
    doc.add_heading("COST & VOLUME BY UNIT", 2)
    u_table = doc.add_table(rows=1, cols=5)
    u_table.style = "Table Grid"
    for i, h in enumerate(["Unit","Line Items","Total Cost","% of Total","Avg Cost/Item"]):
        c = u_table.rows[0].cells[i]
        c.text = h
        set_cell_bg(c, "2E5FA3")
        c.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        c.paragraphs[0].runs[0].font.bold = True
    for unit, v in unit_rows_sorted:
        add_table_row(u_table, [unit, v["items"], f"\xa3{v['cost']:,.2f}", f"{v['cost']/tc*100:.1f}%", f"\xa3{v['cost']/v['items']:.2f}"])
    add_table_row(u_table, ["TOTAL", n, f"\xa3{tc:,.2f}", "100%", f"\xa3{tc/n:.2f}"], bold=True, bg="1F3864")

    # Top Customers
    doc.add_heading("TOP CUSTOMERS BY COST", 2)
    c_table = doc.add_table(rows=1, cols=3)
    c_table.style = "Table Grid"
    for i, h in enumerate(["Customer","Total Cost","% of Total"]):
        c = c_table.rows[0].cells[i]
        c.text = h
        set_cell_bg(c, "2E5FA3")
        c.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        c.paragraphs[0].runs[0].font.bold = True
    for cust, cost in stats["top_customers"]:
        add_table_row(c_table, [cust, f"\xa3{cost:,.2f}", f"{cost/tc*100:.1f}%"])

    # Anomalies
    doc.add_heading("ANOMALIES & DATA QUALITY FLAGS", 2)
    a_table = doc.add_table(rows=1, cols=5)
    a_table.style = "Table Grid"
    for i, h in enumerate(["Flag","Count","% of Total","Severity","Action Required"]):
        c = a_table.rows[0].cells[i]
        c.text = h
        set_cell_bg(c, "C00000")
        c.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        c.paragraphs[0].runs[0].font.bold = True
    sev_bg = {"CRITICAL":"C00000","HIGH":"FF8C00","MEDIUM":"FFD700","RESOLVED":"375623"}
    for flag, count, pct, severity, action in anomalies:
        row = add_table_row(a_table, [flag, count, pct, severity, action])
        set_cell_bg(row.cells[3], sev_bg.get(severity, "999999"))

    if extra_info:
        doc.add_heading("SPECIFIC FOCUS", 2)
        doc.add_paragraph(extra_info)

    doc.add_paragraph(f"\nGenerated by Claude Code  \u2014  {today}")

    word_path = base_name + ".docx"
    doc.save(word_path)
    print(f"Word:  {word_path}")
    subprocess.Popen(["start", word_path], shell=True)

# ── PDF Output ────────────────────────────────────────────
if "pdf" in outputs:
    pdf_path = base_name + ".pdf"
    try:
        import win32com.client
        word_app = win32com.client.Dispatch("Word.Application")
        word_app.Visible = False
        # Use Word to convert docx to PDF if word output was also generated
        if "word" in outputs and os.path.exists(word_path):
            doc_com = word_app.Documents.Open(os.path.abspath(word_path))
            doc_com.SaveAs(os.path.abspath(pdf_path), FileFormat=17)
            doc_com.Close()
        else:
            # Build a temporary word doc just for PDF conversion
            from docx import Document as DocX
            tmp_path = base_name + "_tmp.docx"
            # Re-save the HTML as a basic Word doc for PDF
            tmp_doc = DocX()
            tmp_doc.add_heading("PARTS DATABASE  \u2014  SUMMARY", 0)
            tmp_doc.add_paragraph(f"Date: {today}  |  Sheet: {src_label}  |  {n} items  |  Total: \xa3{tc:,.2f}")
            if extra_info:
                tmp_doc.add_paragraph(f"Focus: {extra_info}")
            tmp_doc.save(tmp_path)
            doc_com = word_app.Documents.Open(os.path.abspath(tmp_path))
            doc_com.SaveAs(os.path.abspath(pdf_path), FileFormat=17)
            doc_com.Close()
            os.remove(tmp_path)
        word_app.Quit()
        print(f"PDF:   {pdf_path}")
        subprocess.Popen(["start", pdf_path], shell=True)
    except Exception as e:
        print(f"PDF failed: {e}")

# ── Summary ───────────────────────────────────────────────
print(f"\nDone | {src_label} | {n} items | \xa3{tc:,.2f} | Outputs: {', '.join(sorted(outputs))}")
